"""
build_detailed_project_plan.py

프로젝트 상세 계획서 DOCX를 생성하는 문서화 스크립트입니다.
실제 웹앱 기능을 실행하지는 않고, 퍼스널컬러 진단/옷장/추천 도메인의 구현 계획과 산출물을 문서로 정리합니다.

발표자료와 보고서에 사용하는 WBS, 단계별 개발 계획, 평가 기준 대응 내용을 문서 형태로 재구성할 때 사용하는 보조 파일입니다.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "퍼스널컬러_AI_옷장_프로젝트_상세계획서.docx"


NAVY = "18304A"
BLUE = "2F6F9F"
TEAL = "2E7D73"
MINT = "DCEFEA"
LIGHT_BLUE = "EAF3FA"
LIGHT_GRAY = "F3F5F7"
PALE_YELLOW = "FFF4CE"
PALE_RED = "FDE7E9"
WHITE = "FFFFFF"


def shade(cell, fill: str) -> None:
    """표 셀 배경색을 지정합니다."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: int = 9) -> None:
    """표 셀의 기존 내용을 지우고 정렬/폰트가 적용된 텍스트를 넣습니다."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 and "\n" not in text else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "맑은 고딕"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def set_borders(table) -> None:
    """Word 표에 전체/내부 테두리를 적용합니다."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D0D7DE")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """계획서의 섹션 제목을 추가합니다."""
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    r.font.name = "맑은 고딕"
    r.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)


def add_body(doc: Document, text: str) -> None:
    """계획서 본문 문단을 추가합니다."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "맑은 고딕"
    r.font.size = Pt(10)


def add_bullets(doc: Document, items: list[str]) -> None:
    """핵심 항목을 bullet list로 추가합니다."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.name = "맑은 고딕"
        r.font.size = Pt(9.5)


def add_note(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    """강조해야 할 평가 기준/핵심 메모를 박스 형태로 추가합니다."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.name = "맑은 고딕"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r2 = p.add_run(body)
    r2.font.name = "맑은 고딕"
    r2.font.size = Pt(9)
    set_borders(table)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    """계획서에 사용할 표를 공통 스타일로 생성합니다."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_borders(table)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, NAVY)
        set_cell_text(cell, h, True, WHITE, 8)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            shade(cells[i], WHITE if len(table.rows) % 2 else LIGHT_GRAY)
            set_cell_text(cells[i], text, False, None, 8)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()


def add_ascii_diagram(doc: Document, title: str, diagram: str, fill: str = "F7FBFF") -> None:
    """아키텍처나 데이터 흐름을 ASCII 박스로 삽입합니다."""
    add_heading(doc, title, 3)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(diagram.strip())
    r.font.name = "Consolas"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("1F2937")
    set_borders(table)
    doc.add_paragraph()


def build() -> None:
    """상세 계획서 DOCX 전체를 구성하고 저장합니다."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "맑은 고딕"
    styles["Normal"].font.size = Pt(10)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "맑은 고딕"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("퍼스널컬러 기반 AI 옷장 시스템")
    r.bold = True
    r.font.name = "맑은 고딕"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("프로젝트 상세계획서")
    r.bold = True
    r.font.name = "맑은 고딕"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    add_note(
        doc,
        "문서 범위",
        "현재 통합 React/Vite 프로젝트와 기존 Word 기술 설계 보고서를 함께 분석하여, 과제 제출용으로 구현 내용, 아키텍처, 흐름도, 모듈별 계획, 평가 계획을 한 문서에 정리했다.",
        MINT,
    )
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["프로젝트명", "퍼스널컬러 기반 AI 옷장 및 코디 추천 시스템"],
            ["기술 스택", "React 19, TypeScript, Vite 6, Tailwind CSS 4, MediaPipe Tasks Vision, Open-Meteo API"],
            ["핵심 기능", "사진 기반 퍼스널컬러 분석, 8문항 설문 보정, 12시즌 결과, 디지털 옷장, 날씨 기반 코디 추천, 저장 코디/가상착용 미리보기"],
            ["저장 방식", "브라우저 localStorage 기반 시제품 데이터 저장"],
            ["작성 기준", "현재 코드 구현 + personal_color_ai_wardrobe_report.docx 설계 방향"],
        ],
        [3, 13],
    )

    add_heading(doc, "1. 프로젝트 개요", 1)
    add_body(
        doc,
        "이 프로젝트는 사용자의 얼굴 사진과 설문 응답을 결합해 12시즌 퍼스널컬러를 산출하고, 사용자가 보유하거나 카탈로그에서 선택한 의류를 디지털 옷장으로 관리한 뒤, 퍼스널컬러 적합도와 날씨 조건을 반영해 코디를 추천하는 웹 애플리케이션이다.",
    )
    add_bullets(
        doc,
        [
            "문제의식: 퍼스널컬러 결과가 실제 옷장 선택으로 연결되지 못하는 문제를 해결한다.",
            "핵심 가치: 새 옷 구매 유도보다 사용자가 이미 보유한 옷을 더 정확하게 활용하도록 돕는다.",
            "서비스 성격: 모바일/데스크톱 브라우저에서 동작하는 프론트엔드 중심 시제품이다.",
            "기존 Word 문서와의 관계: 기존 보고서는 U²-Net, KMeans, 색상 DB 등 확장 설계까지 포함하며, 현재 통합 코드는 퍼스널컬러 분석과 옷장/추천 기능을 우선 구현했다.",
        ],
    )

    add_heading(doc, "2. 구현 현황 요약", 1)
    add_table(
        doc,
        ["영역", "구현 상태", "주요 파일", "설명"],
        [
            ["퍼스널컬러 진단", "구현", "PhotoAnalyzer.tsx, photoAnalysis.ts, geminiService.ts", "MediaPipe 얼굴 랜드마크와 ROI 색상 샘플링, 설문 점수 통합"],
            ["12시즌 팔레트", "구현", "personalColorWorkbook.ts, seasonContent.ts", "12시즌별 traits, 팔레트, 설명, 워스트 컬러 관리"],
            ["디지털 옷장", "구현", "App.tsx", "옷장 생성/수정/삭제, 카탈로그 추가, 수동 의류 등록, 검색/필터"],
            ["날씨 추천", "구현", "useWeather.ts, weather.ts", "Open-Meteo, 대기질 API, 기온 구간별 착장 점수"],
            ["코디 저장/가상착용", "부분 구현", "App.tsx", "추천 코디 저장 및 이미지 미리보기 중심"],
            ["의류 사진 자동 색상 추출", "확장 설계", "기존 Word 보고서", "U²-Net/KMeans 기반 설계는 문서화되어 있으나 현재 통합 코드에는 미구현"],
            ["서버/DB", "미구현", "-", "현재는 localStorage 시제품 구조. 백엔드와 계정 시스템은 확장 과제"],
        ],
        [2.7, 2, 4.2, 7.2],
    )

    add_heading(doc, "3. 전체 아키텍처", 1)
    add_ascii_diagram(
        doc,
        "그림 1. 시스템 계층 구조",
        """
[사용자 브라우저]
  ├─ 카메라/파일 입력
  ├─ 설문 응답
  ├─ 옷장/의류 조작
  └─ 위치 권한
        │
        ▼
[React UI Layer]
  App.tsx ─ PhotoAnalyzer ─ Questionnaire ─ ResultDisplay
        │
        ▼
[Domain Logic Layer]
  photoAnalysis.ts  : ROI 샘플링, 조명 보정, 품질 점수
  geminiService.ts  : 사진 점수, 설문 점수, 하이브리드 융합
  weather.ts        : 날씨 API, 기온 구간, 우산/미세먼지 판단
        │
        ▼
[Data Layer]
  personalColorWorkbook.ts : 12시즌 팔레트/traits
  seasonContent.ts         : 시즌 설명/추천/피해야 할 색
  localStorage             : 진단 결과, 옷장, 의류, 저장 코디
        │
        ▼
[External Services]
  MediaPipe CDN Model / Open-Meteo / Air Quality / Reverse Geocoding
""",
    )
    add_ascii_diagram(
        doc,
        "그림 2. 사용자 이용 흐름",
        """
홈
 ├─ 퍼스널컬러 진단
 │   ├─ 카메라 권한 요청
 │   ├─ 얼굴 인식 및 3초 자동 촬영
 │   ├─ 얼굴 ROI 색상 분석
 │   ├─ 8문항 설문
 │   └─ 12시즌 결과/근거/팔레트 표시
 ├─ 옷장
 │   ├─ 기본 옷장 목록
 │   ├─ 카탈로그 의류 선택
 │   ├─ 수동 의류 등록
 │   └─ 검색/필터/삭제/이름 변경
 ├─ 추천
 │   ├─ 퍼스널컬러 적합도 계산
 │   ├─ 현재 날씨 및 기온 구간 반영
 │   ├─ 목적 모드 선택
 │   └─ 코디 추천 및 저장
 └─ 저장/가상착용/설정
""",
    )

    add_heading(doc, "4. 퍼스널컬러 분석 모듈", 1)
    add_table(
        doc,
        ["단계", "처리 내용", "핵심 구현"],
        [
            ["얼굴 검출", "MediaPipe Face Landmarker로 얼굴 landmark 검출", "GPU delegate 우선, 실패 시 CPU fallback"],
            ["ROI 생성", "볼, 이마, 코 주변, 눈가, 턱선, 홍채, 눈썹, 입술, 헤어라인 영역 생성", "buildSampleRegions()"],
            ["색상 샘플링", "극단 밝기 제거, saturation filter, Lab median medoid 적용", "sampleSkinRegion(), labMedianMedoid()"],
            ["조명 보정", "우측 하단 흰 종이 기준, 중립 배경, 모서리 fallback 순서", "WHITE_REFERENCE_REGION_RATIO, BackgroundCalibration"],
            ["품질 평가", "노출, 좌우 대칭, 얼굴 크기, 배경 중립성, 색상 구분도 평가", "qualityBreakdown"],
            ["특징 벡터", "temperature, lightness, clarity, contrast, mutedScore 정규화", "measureColorFeatures()"],
            ["시즌 점수", "팔레트 Delta E + traits 유사도 + 뮤트/대비 보정", "analyzePhotoColors()"],
        ],
        [2.3, 6.8, 6.5],
    )
    add_ascii_diagram(
        doc,
        "그림 3. 사진 분석 파이프라인",
        """
카메라 프레임
  → MediaPipe 478 landmark
  → 얼굴 bounding box 및 ROI 15종 생성
  → ROI 픽셀 추출
  → 어두운/밝은 오염 픽셀 제거
  → 흰 종이 또는 배경 기반 조명 보정
  → RGB/HSL/Lab 특징 계산
  → 12시즌별 사진 점수 산출
""",
    )
    add_note(
        doc,
        "중요 구현 포인트",
        "현재 구현은 단순히 RGB 평균을 내지 않고, ROI별 목적에 맞게 샘플링 방법을 다르게 둔다. 피부는 Lab median medoid로 안정화하고, 입술은 붉은/핑크 계열 후보를 우선하며, 조명 보정 출처를 결과 데이터에 남긴다.",
        PALE_YELLOW,
    )

    add_heading(doc, "5. 설문 및 하이브리드 융합", 1)
    add_body(doc, "설문은 손목 혈관, 금속 액세서리, 흰색 의류, 햇빛 반응, 비비드 컬러, 뮤트 컬러, 대비 선호, 색의 깊이 등 8문항으로 구성된다. 각 선택지는 temperature, lightness, clarity, contrast 네 축에 가중치를 부여한다.")
    add_table(
        doc,
        ["축", "의미", "추천에 미치는 영향"],
        [
            ["temperature", "웜/쿨 방향", "봄·가을 또는 여름·겨울 계열 우선순위"],
            ["lightness", "라이트/딥 방향", "밝은 파스텔 또는 깊은 다크 톤 추천"],
            ["clarity", "브라이트/뮤트 방향", "선명한 원색 또는 회색기 있는 색 추천"],
            ["contrast", "대비 선호/소화력", "고대비 코디 또는 부드러운 톤온톤 코디"],
        ],
        [3, 5, 8],
    )
    add_ascii_diagram(
        doc,
        "그림 4. 사진+설문 융합 구조",
        """
사진 분석 결과
  ├─ 색상 특징 벡터
  ├─ 시즌별 photoScore
  └─ photoQuality
        │
        ├─ 동적 가중치: photo = clamp(0.22 + photoQuality×0.14, 0.22, 0.36)
        ▼
설문 결과 ── questionnaireScore ──┐
                                  ▼
                   fusedScore = photoScore×photoWeight
                              + questionScore×questionWeight
                                  ▼
                 Top1 시즌 / Top2 시즌 / 경계 여부 / 근거 설명
""",
    )

    add_heading(doc, "6. 옷장 관리 및 추천 엔진", 1)
    add_body(doc, "옷장 기능은 App.tsx 내부 상태와 localStorage를 중심으로 동작한다. 초기 카탈로그와 기본 옷장을 제공하고, 사용자는 카탈로그 추가 또는 수동 등록으로 의류 데이터를 구성한다.")
    add_table(
        doc,
        ["데이터", "주요 필드", "사용 목적"],
        [
            ["Wardrobe", "id, name, createdAt", "옷장 단위 분류"],
            ["ClothingItem", "category, type, color, size, brand, representativeHex, seasonTag, availabilityStatus", "의류 카드와 추천 점수 계산"],
            ["ScoredClothingItem", "personalFitScore, fitGrade, fitReason, avoidRisk", "퍼스널컬러 적합도 표시"],
            ["OutfitRecommendation", "score, personalScore, harmonyScore, weatherScore, stabilityScore, items, reason", "추천 결과 설명"],
            ["SavedOutfit", "itemIds, colorHexes, weatherBand, mode, savedAt", "저장 코디와 가상착용 미리보기"],
        ],
        [3, 7.5, 5.5],
    )
    add_ascii_diagram(
        doc,
        "그림 5. 코디 추천 점수 흐름",
        """
의류 목록
  → 추천 제외/세탁중 필터링
  → 상의×하의 조합 생성
  → 필요 시 아우터/신발 추가
  → personalScore   : 시즌 팔레트와 의류 대표색 Delta E
  → weatherScore    : 기온 구간별 적합 의류 규칙
  → harmonyScore    : 같은 색/뉴트럴 포함 여부 중심
  → stabilityScore  : 모두 보유중이면 가산
  → 최종 점수 = personal 42% + weather 28% + harmony 20% + stability 10%
""",
    )

    add_heading(doc, "7. 날씨 API 및 생활 맥락", 1)
    add_bullets(
        doc,
        [
            "위치 권한이 있으면 현재 좌표를 사용하고, 실패하면 서울 기준 fallback을 사용한다.",
            "Open-Meteo forecast API에서 현재 기온, 체감온도, 날씨 코드, 강수량, 풍속, 일 최고/최저 기온을 가져온다.",
            "Air Quality API에서 PM10, PM2.5, European AQI를 받아 미세먼지 등급과 마스크 권장 여부를 만든다.",
            "기온은 4도 이하, 5~8도, 9~11도, 12~16도, 17~19도, 20~22도, 23~27도, 28도 이상 8개 추천 구간으로 정규화한다.",
        ],
    )

    add_heading(doc, "8. UI/UX 설계", 1)
    add_table(
        doc,
        ["화면", "사용자 목표", "주요 UI 요소"],
        [
            ["홈", "진단 상태, 옷장 상태, 추천 진입을 한눈에 파악", "요약 카드, 바로가기"],
            ["퍼스널컬러", "촬영→설문→결과 확인", "카메라 오버레이, 카운트다운, 설문 카드, 결과 패널"],
            ["옷장", "옷장별 의류 관리", "검색, 카테고리 탭, grid/list, 카탈로그, 수동 등록"],
            ["추천", "오늘 입을 코디 선택", "날씨 카드, 목적 모드, 옷장 선택, 추천 리스트"],
            ["저장/가상착용", "마음에 드는 코디 재확인", "저장 목록, 이미지 미리보기"],
            ["설정", "개인 결과 및 전체 데이터 초기화", "초기화 버튼, 안내"],
        ],
        [2.8, 5.3, 7.9],
    )

    add_heading(doc, "9. 기존 Word 보고서 반영 사항", 1)
    add_table(
        doc,
        ["보고서 설계 내용", "현재 코드 반영", "계획서 반영 방식"],
        [
            ["퍼스널컬러 얼굴 분석", "강하게 반영", "MediaPipe ROI, 조명 보정, 12시즌 점수로 구체화"],
            ["의류 사진 배경 제거", "미반영", "향후 확장 과제로 분리"],
            ["KMeans 대표색 추출", "부분 대체", "현재는 카탈로그/수동 대표색, 향후 자동 추출 계획"],
            ["색상 DB/메타데이터", "부분 반영", "COLOR_META, SEASON_PROFILES, ClothingItem 구조로 반영"],
            ["설명 가능한 추천", "반영", "fitReason, reason, 결과 근거 패널로 구현"],
            ["상황 기반 추천", "반영", "날씨/기온/대기질 기반 추천으로 구현"],
        ],
        [4, 3, 8],
    )

    add_heading(doc, "10. 개발 일정 계획", 1)
    add_table(
        doc,
        ["주차", "목표", "작업 내용", "산출물"],
        [
            ["1주차", "요구사항 및 자료 정리", "기존 보고서, 코드, UI 흐름 분석", "요구사항 명세, 기능 목록"],
            ["2주차", "퍼스널컬러 분석 안정화", "ROI, 조명 보정, 품질 점수, 설문 가중치 검토", "진단 테스트 케이스"],
            ["3주차", "옷장 기능 정리", "카탈로그/수동 등록, localStorage, 검색/필터 확인", "옷장 기능 시연"],
            ["4주차", "추천 엔진 개선", "점수 가중치, 날씨 구간, 추천 이유 문장 보강", "추천 결과 샘플"],
            ["5주차", "반응형 및 접근성", "모바일 카메라, 버튼 크기, 텍스트 가독성, 오류 메시지 개선", "모바일/PC QA표"],
            ["6주차", "검증 및 발표 자료", "빌드, 주요 시나리오 테스트, 한계/개선 정리", "최종 보고서, 발표 자료"],
        ],
        [1.8, 3.2, 7.2, 4.7],
    )

    add_heading(doc, "11. 테스트 및 평가 계획", 1)
    add_table(
        doc,
        ["평가 항목", "방법", "성공 기준"],
        [
            ["빌드 안정성", "npm run build, npm run lint", "타입 오류 및 빌드 실패 없음"],
            ["카메라 진단", "정면/측면/어두운 조명/흰 종이 기준 비교", "품질 점수와 경고가 상황에 맞게 변화"],
            ["설문 융합", "동일 사진+다른 설문 응답 비교", "최종 시즌이 설문 방향을 합리적으로 반영"],
            ["옷장 데이터", "생성, 수정, 삭제, 새로고침 후 유지 확인", "localStorage 데이터 유지"],
            ["추천 품질", "계절별/날씨별 의류 조합 테스트", "추천 이유가 점수 구성과 모순되지 않음"],
            ["반응형", "모바일/데스크톱 뷰포트 확인", "주요 버튼과 카드가 겹치지 않음"],
        ],
        [3, 7, 6],
    )

    add_heading(doc, "12. 한계와 개선 방향", 1)
    add_note(
        doc,
        "현재 한계",
        "현재 통합본은 프론트엔드 시제품이므로 서버 DB, 사용자 계정, 실제 의류 사진 자동 색상 추출, 착용샷 기반 세그멘테이션은 아직 구현 범위 밖이다. 또한 스마트폰 카메라의 자동 보정과 조명 차이는 완전히 제거하기 어렵다.",
        PALE_RED,
    )
    add_bullets(
        doc,
        [
            "1단계 개선: 진단 결과 히스토리와 옷장 데이터를 IndexedDB 또는 백엔드 DB로 이전한다.",
            "2단계 개선: 의류 사진 업로드 후 배경 제거와 대표색 추출을 자동화한다.",
            "3단계 개선: 사용자 피드백을 저장해 추천 가중치를 개인별로 조정한다.",
            "4단계 개선: 일정/장소/목적 데이터를 더 세분화해 출근, 데이트, 발표, 여행 등 맥락별 추천을 강화한다.",
            "5단계 개선: 실제 사용자 평가 데이터를 수집해 퍼스널컬러 결과와 추천 만족도를 정량 검증한다.",
        ],
    )

    add_heading(doc, "13. 결론", 1)
    add_body(
        doc,
        "본 프로젝트는 퍼스널컬러 진단만 따로 제공하는 서비스가 아니라, 진단 결과를 디지털 옷장과 날씨 기반 코디 추천까지 연결한 통합형 시제품이다. 현재 구현은 React 기반 단일 페이지 앱으로 충분히 시연 가능하며, 기존 Word 보고서의 확장 설계를 바탕으로 의류 자동 색상 추출과 서버 기반 데이터 관리까지 발전시킬 수 있다. 과제 관점에서는 기술 구현, 데이터 흐름, UI 흐름, 추천 점수 체계, 평가 계획이 하나의 구조 안에서 연결된다는 점이 핵심 성과다.",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "부록 A. 주요 파일 역할", 1)
    add_table(
        doc,
        ["파일", "역할"],
        [
            ["src/App.tsx", "앱 라우팅, 홈/옷장/추천/저장/설정 화면, localStorage 상태 관리, 추천 엔진"],
            ["src/components/PhotoAnalyzer.tsx", "카메라 실행, 얼굴 추적, 자동 촬영, 분석 진행 UI"],
            ["src/components/Questionnaire.tsx", "8문항 설문 UI와 응답 수집"],
            ["src/components/ResultDisplay.tsx", "퍼스널컬러 결과, 측정 데이터, 개발자 모드 표시"],
            ["src/services/photoAnalysis.ts", "랜드마크 기반 ROI, 색상 샘플링, 조명 보정, 품질 계산"],
            ["src/services/geminiService.ts", "사진/설문 점수 계산과 하이브리드 융합"],
            ["src/services/colorUtils.ts", "RGB, HSL, Lab 변환과 Delta E 계산"],
            ["src/services/faceLandmarker.ts", "MediaPipe Face Landmarker 로딩"],
            ["src/personalColorWorkbook.ts", "12시즌 표준 팔레트와 traits 데이터"],
            ["src/lib/weather.ts", "Open-Meteo 및 대기질 API 처리"],
            ["src/hooks/useWeather.ts", "위치 권한, fallback, 날씨 상태 hook"],
            ["personal_color_ai_wardrobe_report.docx", "기존 기술 설계 통합 보고서"],
        ],
        [5, 11],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
