"""
build_digital_wardrobe_report_docx.py

디지털 옷장과 의류 등록/추천 관계를 설명하는 Word 보고서를 생성하는 문서화 스크립트입니다.
앱 화면 안에 포함되지는 않으며, 발표/보고서 제출용 산출물을 자동 생성하기 위해 사용합니다.

코드 관점에서는 프로젝트의 옷장 도메인, 의류 데이터 모델, 추천 알고리즘 연결 구조를 문서 형식으로 풀어내는 보조 도구입니다.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "디지털 옷장과 의류 등록 모듈 및 추천 알고리즘 관계.md"
OUTPUT = ROOT / "디지털 옷장과 의류 등록 모듈 및 추천 알고리즘 관계.docx"

ACCENT = "1F4E79"
LIGHT = "EAF2FB"
HEADER = "D8E8F7"
TEXT = RGBColor(31, 41, 55)


def set_cell_shading(cell, fill: str) -> None:
    """표 셀 배경색을 설정합니다."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    """Word 표 셀 내부 여백을 설정합니다."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="CAD6E2", size="6") -> None:
    """Word 표 테두리 색상과 두께를 설정합니다."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_run(run, bold=False, size=10.5, color: RGBColor | None = None) -> None:
    """문서 전체에서 사용하는 글꼴/크기/색상 규칙을 run에 적용합니다."""
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color or TEXT


def add_para(doc: Document, text: str = "", style: str | None = None, space_after=3, align=None):
    """본문 문단을 추가하고 기본 줄간격/여백을 설정합니다."""
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.08
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        style_run(run)
    return paragraph


def add_heading(doc: Document, text: str, level: int) -> None:
    """보고서 섹션 제목을 추가합니다."""
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {min(level, 3)}"
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    style_run(run, bold=True, size={1: 17, 2: 13.5, 3: 11.5}.get(level, 11))
    run.font.color.rgb = RGBColor(31, 78, 121) if level <= 2 else RGBColor(55, 65, 81)


def add_callout(doc: Document, title: str, body: str) -> None:
    """핵심 요약이나 주의점을 강조 박스로 추가합니다."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="B6CCE2")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    style_run(run, bold=True, size=10.5, color=RGBColor(31, 78, 121))
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    run2 = p2.add_run(body)
    style_run(run2, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    """마크다운에서 읽은 표 데이터를 Word 표로 변환합니다."""
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for idx, value in enumerate(rows[0]):
        cell = header_cells[idx]
        set_cell_shading(cell, HEADER)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        style_run(run, bold=True, size=9)
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 14 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            style_run(run, size=8.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def parse_markdown_tables(lines: list[str], start: int):
    """마크다운 표 시작 위치에서 행 데이터를 파싱하고 다음 읽기 위치를 반환합니다."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_code_block(doc: Document, code: list[str]) -> None:
    """코드/구조 설명 블록을 monospace 박스로 추가합니다."""
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="D1D5DB")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")
    set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.05
    text = "\n".join(code)
    run = paragraph.add_run(text[:1800])
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.2)
    run.font.color.rgb = RGBColor(17, 24, 39)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_document() -> None:
    """마크다운 보고서를 읽어 표, 제목, 코드 블록이 포함된 DOCX로 변환합니다."""
    md = SOURCE.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(16)
    run = title.add_run("디지털 옷장과 의류 등록 모듈\n및 추천 알고리즘 관계")
    style_run(run, bold=True, size=22, color=RGBColor(31, 78, 121))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(38)
    run = subtitle.add_run("4.3 모듈 상세분석 보고서")
    style_run(run, bold=True, size=13, color=RGBColor(75, 85, 99))

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta, color="B6CCE2")
    metadata = [
        ("문서 구분", "상세분석 보고서"),
        ("대상 시스템", "퍼스널컬러 기반 AI 디지털 옷장"),
        ("분석 범위", "의류 등록, 저장 구조, 추천 알고리즘 연계"),
        ("작성 기준", "현재 React/Vite 구현 코드 및 로컬 데이터 구조"),
    ]
    for row, (k, v) in zip(meta.rows, metadata):
        set_cell_shading(row.cells[0], HEADER)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = row.cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p1.add_run(k), bold=True, size=9.5, color=RGBColor(31, 78, 121))
        p2 = row.cells[1].paragraphs[0]
        style_run(p2.add_run(v), size=9.5)

    doc.add_page_break()
    add_heading(doc, "보고서 요약", 1)
    add_callout(
        doc,
        "핵심 요지",
        "디지털 옷장 모듈은 단순 이미지 보관함이 아니라 추천 알고리즘의 입력 데이터베이스이다. "
        "의류 등록 시 카테고리, 타입, 대표 색상, 계절 태그, 보유 상태, 원본 카탈로그 ID가 함께 저장되며, "
        "추천 모듈은 이 정보를 바탕으로 기온 필터, 퍼스널컬러 적합도, 색상 조화, 보유 상태를 계산한다.",
    )
    add_table(
        doc,
        [
            ["구분", "역할", "추천 알고리즘과의 관계"],
            ["옷장", "의류 묶음 단위", "선택된 옷장만 추천 후보로 사용"],
            ["의류 아이템", "추천의 최소 데이터 단위", "카테고리/타입/색상/상태가 추천 기준으로 활용"],
            ["카탈로그", "관리자 준비 의류 DB", "사용자 옷장에 들어가면 ClothingItem으로 변환"],
            ["저장 코디", "추천 결과 저장값", "아이템 ID 기반으로 재조회하여 홈/가상착용에 표시"],
        ],
    )

    i = 0
    in_code = False
    code_buf: list[str] = []
    skip_first = True
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if skip_first and line.startswith("# "):
            skip_first = False
            i += 1
            continue
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(raw)
            i += 1
            continue
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_markdown_tables(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            text = line.lstrip("#").strip()
            add_heading(doc, text, level)
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(line[2:])
            style_run(run, size=9.8)
        elif re.match(r"^\d+\. ", line):
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(re.sub(r"^\d+\. ", "", line))
            style_run(run, size=9.8)
        else:
            add_para(doc, line, space_after=5)
        i += 1

    doc.core_properties.title = "디지털 옷장과 의류 등록 모듈 및 추천 알고리즘 관계"
    doc.core_properties.subject = "4.3 디지털 옷장과 의류 등록 모듈 상세분석 보고서"
    doc.core_properties.author = "Fitly Project"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
