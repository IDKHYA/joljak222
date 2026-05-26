"""
build_assignment_format_plan.py

과제 양식에 맞춘 프로젝트 계획서 DOCX를 생성하는 문서화 스크립트입니다.
앱 본체의 실행 로직과 분리되어 있으며, 제출 문서의 형식과 본문 구성을 자동화하는 목적입니다.

프로젝트의 문제 정의, 목표, 기술 스택, 추진 일정, 기대 효과를 과제 제출 형식에 맞게 정리하는 보조 도구로 보면 됩니다.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "퍼스널컬러_AI_옷장_과제양식_상세계획서.docx"

NAVY = "18304A"
BLUE = "2F6F9F"
TEAL = "2E7D73"
WHITE = "FFFFFF"
LIGHT_GRAY = "F3F5F7"
LIGHT_BLUE = "EAF3FA"
MINT = "DCEFEA"
YELLOW = "FFF4CE"
RED = "FDE7E9"


def shade(cell, fill: str) -> None:
    """표 셀 배경색을 지정해 문서 구역을 시각적으로 구분합니다."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(table) -> None:
    """python-docx 표에 기본 테두리를 적용합니다."""
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "D0D7DE")


def run_text(paragraph, text: str, size: int = 10, bold: bool = False, color: str | None = None, font: str = "맑은 고딕"):
    """문단에 글자를 추가하면서 폰트, 크기, 굵기, 색상을 한 번에 설정합니다."""
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def heading(doc: Document, text: str, level: int) -> None:
    """문서 제목 계층을 통일된 색상/크기로 추가합니다."""
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run_text(p, text, 16 if level == 1 else 13 if level == 2 else 11, True, NAVY if level == 1 else BLUE)


def body(doc: Document, text: str) -> None:
    """일반 본문 문단을 추가합니다."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    run_text(p, text, 10)


def bullets(doc: Document, items: list[str]) -> None:
    """문자열 목록을 Word bullet list로 추가합니다."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run_text(p, item, 9)


def note(doc: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    """강조 박스 형태의 메모를 추가합니다."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    run_text(p, title + "\n", 10, True, NAVY)
    run_text(p, text, 9)
    borders(table)
    doc.add_paragraph()


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    """헤더와 행 데이터를 받아 제출 문서용 표를 생성합니다."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    borders(t)
    for i, header in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_text(p, header, 8, True, WHITE)
    for ridx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for i, text in enumerate(row):
            c = cells[i]
            shade(c, WHITE if ridx % 2 else LIGHT_GRAY)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 16 and "\n" not in text else WD_ALIGN_PARAGRAPH.LEFT
            run_text(p, text, 8)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()


def diagram(doc: Document, title: str, text: str) -> None:
    """ASCII 구조도를 monospace 박스로 삽입합니다."""
    heading(doc, title, 3)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, "F7FBFF")
    p = c.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    run_text(p, text.strip(), 8, False, "1F2937", "Consolas")
    borders(t)
    doc.add_paragraph()


def progress_bar(percent: int) -> str:
    """퍼센트 값을 10칸 진행 막대 문자열로 바꿉니다."""
    full = round(percent / 10)
    return "■" * full + "□" * (10 - full) + f" {percent}%"


def build() -> None:
    """과제 양식에 맞춘 전체 DOCX 문서를 처음부터 끝까지 생성합니다."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.name = "맑은 고딕"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_text(p, "퍼스널컬러 기반 AI 옷장 시스템", 23, True, NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_text(p, "과제 양식 기반 상세계획서", 17, True, BLUE)
    note(
        doc,
        "작성 기준",
        "현재 통합_퍼컬_옷장 React/Vite 프로젝트와 기존 personal_color_ai_wardrobe_report.docx를 함께 검토하여 과제 제출 목차에 맞게 재구성했다.",
        MINT,
    )

    heading(doc, "1. 서론", 1)
    heading(doc, "1.1 과제 배경 및 필요성", 2)
    body(doc, "현대 사용자는 퍼스널컬러 진단 결과를 알고 있어도 실제 보유 의류와 연결해 매일 입을 옷을 고르는 데 어려움을 겪는다. 기존 퍼스널컬러 서비스는 주로 진단 결과와 추천 색상 안내에 머무르고, 디지털 옷장 서비스는 의류 관리에 집중하지만 사용자의 피부톤·계절톤과 현재 날씨까지 함께 고려하는 경우가 제한적이다.")
    bullets(doc, [
        "추진 이유: 퍼스널컬러 진단 결과를 실제 옷장 데이터와 연결해 실사용 가능한 코디 추천으로 확장하기 위함이다.",
        "기존 서비스와 차별성: 얼굴 사진 분석, 설문 보정, 12시즌 팔레트, 보유 의류 관리, 날씨 기반 추천을 하나의 프론트엔드 앱 안에서 연결한다.",
        "사용자 가치: 새 옷 구매보다 이미 가진 옷을 더 정확하게 활용하게 하며, 추천 이유를 함께 제공해 결과 신뢰도를 높인다.",
        "과제 적합성: 컴퓨터 비전, 색채 분석, UI/UX, 추천 알고리즘, 외부 API 활용을 통합적으로 보여 줄 수 있다.",
    ])
    heading(doc, "1.2 과제 목표", 2)
    table(doc, ["목표", "주요 기능", "성능/완성 기준"], [
        ["퍼스널컬러 진단", "카메라 얼굴 인식, ROI 색상 추출, 8문항 설문, 12시즌 결과", "사진 품질 점수 산출, Top1/Top2 시즌 및 신뢰도 표시"],
        ["디지털 옷장", "옷장 생성/수정/삭제, 카탈로그 의류 추가, 수동 의류 등록", "새로고침 후 localStorage 데이터 유지"],
        ["코디 추천", "퍼스널컬러 적합도, 날씨 적합도, 조화도, 안정성 점수 결합", "추천 이유와 점수 구성 표시"],
        ["생활 맥락 반영", "현재 위치 또는 서울 fallback 날씨, 기온 구간, 우산/미세먼지 안내", "Open-Meteo API 응답 기반 날씨 카드 제공"],
        ["시연 완성도", "반응형 UI, 모바일 카메라 대응, 빌드 가능 상태", "npm run lint/build 통과"],
    ], [3, 7, 6])

    heading(doc, "2. 설계 및 구현 현황", 1)
    heading(doc, "2.1 시스템 아키텍처", 2)
    diagram(doc, "그림 1. 전체 시스템 구조도", """
[사용자]
  ├─ 얼굴 촬영/사진 입력
  ├─ 8문항 설문 응답
  ├─ 옷장/의류 관리
  └─ 위치 권한 제공
        │
        ▼
[React UI]
  App.tsx
  ├─ PhotoAnalyzer.tsx
  ├─ Questionnaire.tsx
  └─ ResultDisplay.tsx
        │
        ▼
[분석/추천 로직]
  photoAnalysis.ts  : 얼굴 ROI, 색상 샘플링, 조명 보정
  geminiService.ts  : 사진 점수 + 설문 점수 융합
  weather.ts        : 날씨 API, 기온 구간, 대기질 처리
        │
        ▼
[데이터]
  12시즌 팔레트/traits, 시즌 설명, 의류 메타데이터, localStorage
        │
        ▼
[외부 서비스]
  MediaPipe Face Landmarker / Open-Meteo / Air Quality API
""")
    diagram(doc, "그림 2. 데이터 흐름도", """
카메라 프레임
 → MediaPipe 얼굴 랜드마크
 → 볼/이마/눈/입술/헤어 ROI 생성
 → RGB·HSL·Lab 특징 계산
 → 사진 기반 12시즌 점수

설문 응답
 → temperature/lightness/clarity/contrast 축 점수화
 → 설문 기반 12시즌 점수

사진 점수 + 설문 점수
 → 동적 가중치 융합
 → 최종 퍼스널컬러 결과 저장
 → 의류 대표색과 비교
 → 날씨 조건 반영
 → 코디 추천 결과 출력
""")

    heading(doc, "2.2 세부 설계 사항", 2)
    table(doc, ["구분", "설계 내용", "현재 적용 방식"], [
        ["DB 설계", "시제품 단계에서는 서버 DB 대신 localStorage를 사용", "personalColor, personalHistory, wardrobes, clothing, savedOutfits 키로 저장"],
        ["데이터 모델", "Wardrobe, ClothingItem, ScoredClothingItem, OutfitRecommendation, SavedOutfit", "src/App.tsx와 src/types.ts에 TypeScript 타입으로 정의"],
        ["알고리즘", "사진 분석 → 설문 보정 → 시즌 점수 융합 → 의류 적합도 계산 → 코디 추천", "photoAnalysis.ts, geminiService.ts, App.tsx에 분리 구현"],
        ["UI/UX", "홈, 퍼스널컬러, 옷장, 추천, 저장, 가상착용, 설정 화면", "좌측 사이드바와 모바일 하단 내비게이션 구조"],
        ["확장 설계", "의류 사진 배경 제거, KMeans 대표색 추출, 서버 DB", "기존 Word 보고서에는 설계되어 있으나 현재 코드는 카탈로그/수동 등록 중심"],
    ], [2.5, 7.8, 5.7])
    diagram(doc, "그림 3. 알고리즘 플로우차트", """
시작
  ↓
카메라 권한 확인
  ↓
얼굴 검출 성공?
  ├─ 아니오 → 재촬영 안내
  └─ 예
      ↓
    3초 자동 촬영
      ↓
    ROI 색상 추출 및 조명 보정
      ↓
    사진 점수 계산
      ↓
    설문 8문항 응답
      ↓
    사진/설문 점수 융합
      ↓
    12시즌 Top1/Top2 결과 출력
      ↓
    옷장 의류와 색상 적합도 비교
      ↓
    날씨·목적 반영 코디 추천
      ↓
종료
""")
    diagram(doc, "그림 4. UI/UX 와이어프레임", """
┌─────────────────────────────────────────────┐
│ 사이드바/모바일 내비게이션                  │
├───────────────┬─────────────────────────────┤
│ 홈            │ 진단 결과 요약 / 옷장 현황  │
│ 퍼스널컬러    │ 카메라 → 설문 → 결과        │
│ 옷장          │ 목록 / 상세 / 카탈로그 / 등록│
│ 추천          │ 날씨 카드 / 목적 / 추천 목록│
│ 저장          │ 저장 코디 목록              │
│ 가상착용      │ 저장 코디 이미지 미리보기   │
│ 설정          │ 데이터 초기화               │
└───────────────┴─────────────────────────────┘
""")

    heading(doc, "2.3 현재 구현 범위", 2)
    table(doc, ["모듈", "완료 수준", "주요 파일", "설명"], [
        ["퍼스널컬러 촬영", "완료", "PhotoAnalyzer.tsx", "카메라, 얼굴 추적, 자동 촬영, 오류 안내"],
        ["색상 분석", "완료", "photoAnalysis.ts, colorUtils.ts", "ROI 샘플링, Lab/HSL 변환, 조명 보정, 품질 점수"],
        ["결과 융합", "완료", "geminiService.ts", "사진 점수와 설문 점수의 동적 가중치 결합"],
        ["12시즌 콘텐츠", "완료", "personalColorWorkbook.ts, seasonContent.ts", "시즌별 traits, 팔레트, 추천/피해야 할 색"],
        ["옷장 관리", "완료", "App.tsx", "옷장 CRUD, 카탈로그 추가, 수동 등록, 필터/검색"],
        ["날씨 추천", "완료", "weather.ts, useWeather.ts", "현재 위치/fallback, 기온 구간, 대기질, 우산 안내"],
        ["저장/가상착용", "부분 완료", "App.tsx", "저장 코디와 이미지 미리보기 중심"],
        ["의류 자동 색상 추출", "미구현", "설계 문서", "U²-Net/KMeans 방식은 향후 구현 예정"],
    ], [2.5, 2, 4, 7.5])
    heading(doc, "2.4 주요 소스코드/설계안 설명", 2)
    table(doc, ["파일", "핵심 로직", "기술적 의미"], [
        ["photoAnalysis.ts", "buildSampleRegions, sampleSkinRegion, analyzeFaceSnapshotColors", "얼굴 landmark를 색상 분석 가능한 ROI 데이터로 변환"],
        ["geminiService.ts", "calculateQuestionnaireScores, analyzePhotoColors, fuseResults", "사진 점수와 설문 점수를 하나의 최종 시즌 결과로 통합"],
        ["App.tsx", "scoreItemForPersonalColor, buildRecommendations", "의류 대표색과 시즌 팔레트를 비교하고 코디 점수를 계산"],
        ["weather.ts", "fetchCurrentWeather, getWeatherBandFromTemperature", "외부 날씨 데이터를 추천 가능한 기온 구간으로 정규화"],
        ["personalColorWorkbook.ts", "SEASON_PROFILES", "12시즌별 traits와 팔레트를 코드 데이터로 관리"],
    ], [4, 6, 6])
    note(doc, "핵심 수식", "최종 코디 점수는 personalScore 42% + weatherScore 28% + harmonyScore 20% + stabilityScore 10% 구조로 계산된다. 퍼스널컬러 진단은 사진 품질에 따라 사진 가중치를 22~36% 범위에서 동적으로 조정한다.", YELLOW)

    heading(doc, "3. 진행도 분석", 1)
    heading(doc, "3.1 추진 일정 대비 실적", 2)
    table(doc, ["작업", "초기 계획", "현재 실적", "달성도"], [
        ["프로젝트 구조 및 UI", "1주차 완료", "홈/내비게이션/주요 화면 구성 완료", progress_bar(95)],
        ["퍼스널컬러 사진 분석", "2~3주차 완료", "MediaPipe, ROI, 조명 보정, 결과 화면 완료", progress_bar(90)],
        ["설문 및 결과 융합", "3주차 완료", "8문항 설문과 하이브리드 융합 완료", progress_bar(95)],
        ["옷장 관리", "4주차 완료", "옷장 CRUD, 카탈로그, 수동 등록 완료", progress_bar(85)],
        ["날씨 기반 추천", "5주차 완료", "Open-Meteo/API 연동 및 기온 구간 추천 완료", progress_bar(85)],
        ["가상착용/최종 고도화", "6주차 진행", "저장 코디 미리보기 수준", progress_bar(45)],
        ["의류 자동 색상 추출", "확장 과제", "설계만 존재", progress_bar(20)],
    ], [3, 4, 5, 4])
    diagram(doc, "그림 5. 간트 차트 요약", """
작업                         1주  2주  3주  4주  5주  6주
요구사항/자료 분석            ███
UI 구조 설계                  ███  █
퍼스널컬러 사진 분석              ███  ███
설문/융합 로직                       ███
옷장 관리 기능                           ███  █
날씨 API 및 추천                              ███
반응형/QA/발표 준비                              ███
의류 자동 색상 추출                       설계 → 향후 구현
""")
    heading(doc, "3.2 개발 환경 및 도구 활용", 2)
    table(doc, ["구분", "사용 도구", "활용 내용"], [
        ["프레임워크", "React 19, TypeScript, Vite 6", "SPA 구현, 타입 안정성, 빠른 개발 서버/빌드"],
        ["스타일/UI", "Tailwind CSS 4, shadcn 계열 UI, lucide-react", "반응형 UI, 카드/버튼/아이콘 구성"],
        ["컴퓨터 비전", "@mediapipe/tasks-vision", "Face Landmarker로 얼굴 landmark 검출"],
        ["색상 처리", "Canvas API, RGB/HSL/Lab, Delta E", "피부/머리/눈/입술 색상 특징 계산"],
        ["외부 API", "Open-Meteo, Air Quality API", "현재 날씨, 대기질, 강수 정보 활용"],
        ["상태/저장", "React state, localStorage", "진단 결과와 옷장 데이터 보존"],
        ["협업/버전관리", "Git", "프로젝트 변경 이력 관리"],
        ["검증", "npm run lint, npm run build", "TypeScript 검사와 프로덕션 빌드 확인"],
    ], [3, 4.5, 8.5])
    heading(doc, "3.3 중간 결과물", 2)
    table(doc, ["중간 산출물", "내용", "확인 위치"], [
        ["실행 가능한 웹앱", "Vite 기반 React 앱", "npm run dev 후 localhost:3000"],
        ["퍼스널컬러 진단 화면", "카메라, 얼굴 검출, 자동 촬영, 설문, 결과", "src/components/PhotoAnalyzer.tsx"],
        ["결과 상세 화면", "12시즌, 추천 팔레트, 근거, 개발자 모드", "src/components/ResultDisplay.tsx"],
        ["디지털 옷장 화면", "옷장 목록, 카탈로그 선택, 수동 등록", "src/App.tsx"],
        ["추천 화면", "날씨 카드, 옷장 선택, 추천 코디 저장", "src/App.tsx"],
        ["기술 설계 보고서", "기존 Word 기술 설계 문서", "personal_color_ai_wardrobe_report.docx"],
    ], [4, 7, 5])
    note(doc, "프로토타입 이미지 관련", "현재 문서에는 실행 화면 대신 와이어프레임과 구조도를 포함했다. 발표용 최종본에서는 localhost 실행 후 홈/진단/옷장/추천 화면 캡처를 추가하면 가시성이 더 좋아진다.", LIGHT_BLUE)

    heading(doc, "4. 문제점 및 해결 방안", 1)
    heading(doc, "4.1 기술적 애로사항", 2)
    table(doc, ["문제", "원인", "영향"], [
        ["카메라 색상 왜곡", "스마트폰/브라우저 자동 보정, 조명 차이", "피부색 온도감이 실제와 다르게 계산될 수 있음"],
        ["얼굴 ROI 오염", "머리카락, 안경, 그림자, 메이크업", "피부/눈/입술 대표색 정확도 저하"],
        ["모바일 성능", "MediaPipe 모델 로딩과 실시간 검출 비용", "저사양 기기에서 검출 지연 가능"],
        ["의류 자동 분석 미구현", "배경 제거와 색상 군집화는 추가 구현 필요", "현재는 카탈로그/수동 등록 중심"],
        ["데이터 영속성 한계", "localStorage는 브라우저 단위 저장", "기기 변경/계정 동기화 불가"],
        ["추천 다양성 한계", "현재 조화도 계산이 단순 규칙 중심", "개인 취향 학습과 복잡한 패턴 판단은 부족"],
    ], [3.5, 5, 7.5])
    heading(doc, "4.2 해결 과정 및 계획", 2)
    table(doc, ["이슈", "현재 해결 방법", "향후 대안"], [
        ["색상 왜곡", "흰 종이 기준 영역, 중립 배경, 모서리 fallback 조명 보정", "색상 보정 카드 또는 다중 촬영 평균 적용"],
        ["ROI 오염", "밝기 trimming, saturation filter, Lab median medoid", "얼굴 부위별 segmentation 또는 오염 픽셀 분류 추가"],
        ["모바일 성능", "검출 간격을 데스크톱/모바일로 분리", "모델 lazy loading, 해상도 자동 조절"],
        ["의류 자동 분석", "수동 등록과 카탈로그 방식으로 우선 시연", "U²-Net/GrabCut/KMeans 기반 자동 대표색 추출 구현"],
        ["저장 한계", "localStorage로 시제품 저장", "IndexedDB 또는 Supabase/Firebase 백엔드 도입"],
        ["추천 고도화", "퍼스널컬러/날씨/조화/안정성 가중치 규칙", "사용자 선택 이력 기반 개인화 가중치 학습"],
    ], [3.3, 6.7, 6])
    heading(doc, "4.3 설계 변경 사항", 2)
    table(doc, ["초기 계획", "변경 내용", "변경 사유"], [
        ["의류 사진 자동 색상 추출 즉시 구현", "현재는 카탈로그/수동 등록 중심으로 변경", "시연 안정성과 개발 기간을 우선 고려"],
        ["서버 DB 기반 저장", "localStorage 기반 프론트엔드 시제품으로 변경", "과제 중간 결과물에서는 빠른 시연과 설치 단순성이 중요"],
        ["4계절 중심 진단", "12시즌 세부 진단으로 확장", "라이트/소프트/브라이트/다크 등 세부 톤을 보여야 결과 설득력이 높음"],
        ["날씨는 단순 기온 표시", "기온 구간, 우산, 미세먼지까지 반영", "코디 추천의 실사용성을 높이기 위함"],
        ["단순 결과 화면", "근거/측정 데이터/개발자 모드 포함", "과제 평가에서 알고리즘 설명 가능성을 확보하기 위함"],
    ], [4, 5, 7])

    heading(doc, "5. 향후 추진 계획", 1)
    heading(doc, "5.1 잔여 과업 목록", 2)
    table(doc, ["우선순위", "잔여 과업", "완료 기준"], [
        ["상", "실행 화면 캡처 및 발표용 이미지 추가", "홈/진단/결과/옷장/추천 화면 이미지 확보"],
        ["상", "모바일 카메라 QA", "권한 거부, 검출 실패, 자동 촬영 흐름 검증"],
        ["상", "추천 결과 문장 보강", "점수와 추천 이유가 더 자연스럽게 연결"],
        ["중", "의류 자동 색상 추출 1차 구현", "업로드 이미지에서 대표색 후보 2~3개 제시"],
        ["중", "IndexedDB 또는 백엔드 저장 검토", "데이터 구조 마이그레이션 설계"],
        ["중", "가상착용 화면 고도화", "저장 코디를 더 보기 좋은 레이아웃으로 표시"],
        ["하", "사용자 피드백 기반 추천 개인화", "좋아요/별점 기반 가중치 조정"],
    ], [2.5, 7.5, 6])
    heading(doc, "5.2 최종 일정 계획", 2)
    table(doc, ["주차", "상세 일정", "산출물"], [
        ["1주차", "현재 기능 안정화, lint/build 정기 확인, 주요 버그 정리", "QA 체크리스트"],
        ["2주차", "모바일/데스크톱 화면 캡처, 보고서 이미지 삽입, UI 문구 정리", "중간 결과물 이미지"],
        ["3주차", "추천 이유 문장과 점수 설명 고도화, 날씨 예외 처리 보강", "추천 시연 시나리오"],
        ["4주차", "의류 자동 대표색 추출 최소 기능 검토 또는 확장 설계 상세화", "확장 모듈 설계안"],
        ["5주차", "발표용 데이터 세트 준비, 시연 동선 리허설", "발표 시연 스크립트"],
        ["6주차", "최종 보고서/발표자료 정리, 졸업전시 또는 심사 대비", "최종 제출물"],
    ], [2.5, 9, 4.5])
    heading(doc, "5.3 기대 효과 및 활용 방안", 2)
    bullets(doc, [
        "사용자 측면: 퍼스널컬러 결과를 실제 보유 의류 선택으로 연결해 매일의 코디 결정 시간을 줄인다.",
        "교육/과제 측면: 컴퓨터 비전, 색상공간, 프론트엔드, API, 추천 알고리즘을 통합한 구현 사례로 활용할 수 있다.",
        "서비스 확장 측면: 쇼핑몰 추천, 옷장 재정리, 계절별 캡슐 옷장, 스타일 상담 도구로 확장 가능하다.",
        "데이터 활용 측면: 사용자가 선택한 코디와 만족도 데이터를 축적하면 개인화 추천 정확도를 높일 수 있다.",
        "전시 활용 측면: 카메라 촬영→결과→옷장 추천까지 흐름이 명확해 관람자가 짧은 시간 안에 기능을 이해할 수 있다.",
    ])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
